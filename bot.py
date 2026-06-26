"""
AI Telegram Bot — тарифы, файлы, фото, стриминг ответов.
Стек: Python 3.10+, aiogram 3, Groq API + Google Gemini (OpenAI-совместимый), SQLite.

Премиум работает через Gemini (бесплатный тариф Google AI Studio):
  • модель gemini-2.5-flash (без карты, ~1500 запросов/день),
  • один легальный ключ (или ключ юзера); фолбэк на Groq, если ключа нет,
  • стриминг ответа, Markdown→HTML, разбивка длинных сообщений,
  • фото анализируются через vision-модель Groq.
"""

import asyncio
import logging
import sqlite3
import io
import json
import time
import html
import re
import base64
from datetime import datetime, timedelta

from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart, Command, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.exceptions import TelegramBadRequest, TelegramRetryAfter
from aiogram.types import (
    Message, CallbackQuery,
    InlineKeyboardMarkup, InlineKeyboardButton,
    LabeledPrice, PreCheckoutQuery
)
import httpx

from config import (
    BOT_TOKEN, GROQ_API_KEY, BOT_USERNAME, ADMIN_ID,
    FREE_REQUESTS_PER_DAY, SUBSCRIPTION_DAYS, REFERRAL_BONUS_DAYS,
    GEMINI_API_KEY, GEMINI_BASE_URL, GEMINI_PREMIUM_MODEL, GEMINI_FAST_MODEL,
    DB_PATH,
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger(__name__)

GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
GEMINI_NATIVE_BASE = GEMINI_BASE_URL.rstrip("/")  # native эндпоинт Gemini

# Провайдеры в OpenAI-совместимом формате (Gemini обрабатывается отдельно — native).
ENDPOINTS = {
    "groq": (GROQ_URL, GROQ_API_KEY),
}

# Модель для анализа фото (vision) — через Groq.
GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
# Распознавание голоса (Speech-to-Text) — через Groq, бесплатно.
GROQ_WHISPER_MODEL = "whisper-large-v3-turbo"
# Резервная текстовая модель, если ключа Gemini нет вообще.
GROQ_FALLBACK_MODEL = "llama-3.3-70b-versatile"

TG_LIMIT = 4096
SOFT_LIMIT = 3800
EDIT_THROTTLE = 1.4
MAX_TOKENS = 2048

# Память диалога
MAX_HISTORY = 12        # сколько последних сообщений помнить (user + assistant вместе)
MAX_HIST_CHARS = 4000   # макс длина одного сообщения, сохраняемого в историю

# ─── Тарифы ──────────────────────────────────────────────────────────────────
PLANS = {
    "free": {
        "name": "🆓 Бесплатный", "provider": "groq",
        "model": "llama-3.1-8b-instant",
        "vision": False, "files": False,
        "price_stars": 0, "desc": "3 запроса в день · только текст",
    },
    "basic": {
        "name": "⚡ Базовый", "provider": "groq",
        "model": "llama-3.3-70b-versatile",
        "vision": False, "files": True,
        "price_stars": 150, "desc": "Безлимит · текст + файлы",
    },
    "standard": {
        "name": "🔥 Стандарт", "provider": "groq",
        "model": GROQ_VISION_MODEL,
        "vision": True, "files": True,
        "price_stars": 250, "desc": "Безлимит · файлы + фото",
    },
    "premium": {
        "name": "👑 Премиум", "provider": "gemini",
        "model": GEMINI_PREMIUM_MODEL,
        "vision": True, "files": True,
        "price_stars": 450, "desc": "Gemini 2.5 Flash · всё включено",
    },
}

USER_COLUMNS = [
    "user_id", "username", "plan", "sub_until", "req_today", "req_date",
    "referred_by", "referral_count", "ds_key", "ds_model",
]

class ApiError(Exception):
    def __init__(self, status: int, body: str = ""):
        self.status = status
        self.body = body
        super().__init__(f"API {status}: {body[:200]}")


# ─── База данных ─────────────────────────────────────────────────────────────
def init_db():
    con = sqlite3.connect(DB_PATH)
    con.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            user_id        INTEGER PRIMARY KEY,
            username       TEXT,
            plan           TEXT DEFAULT 'free',
            sub_until      TEXT,
            req_today      INTEGER DEFAULT 0,
            req_date       TEXT,
            referred_by    INTEGER DEFAULT NULL,
            referral_count INTEGER DEFAULT 0
        );
    """)
    existing = {row[1] for row in con.execute("PRAGMA table_info(users)").fetchall()}
    # авто-миграция: добавляем любые недостающие колонки (для старых баз)
    migrations = {
        "plan": "TEXT DEFAULT 'free'",
        "sub_until": "TEXT",
        "req_today": "INTEGER DEFAULT 0",
        "req_date": "TEXT",
        "referred_by": "INTEGER DEFAULT NULL",
        "referral_count": "INTEGER DEFAULT 0",
        "ds_key": "TEXT",
        "ds_model": "TEXT",
    }
    for col, decl in migrations.items():
        if col not in existing:
            con.execute(f"ALTER TABLE users ADD COLUMN {col} {decl}")
    # одноразовая чистка «хвостов» старой OpenModel-версии
    con.execute("UPDATE users SET ds_model=NULL WHERE ds_model IS NOT NULL AND ds_model NOT LIKE 'gemini%'")
    con.execute("UPDATE users SET ds_key=NULL WHERE ds_key IS NOT NULL "
                "AND ds_key NOT LIKE 'AIza%' AND ds_key NOT LIKE 'AQ.%'")
    con.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id      INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            role    TEXT,
            content TEXT,
            ts      TEXT
        );
    """)
    con.execute("CREATE INDEX IF NOT EXISTS idx_history_user ON history(user_id, id)")
    con.commit()
    con.close()


def get_user(user_id: int):
    con = sqlite3.connect(DB_PATH)
    row = con.execute(
        f"SELECT {', '.join(USER_COLUMNS)} FROM users WHERE user_id=?",
        (user_id,)
    ).fetchone()
    con.close()
    return dict(zip(USER_COLUMNS, row)) if row else None


def upsert_user(user_id: int, username: str, referred_by: int = None) -> bool:
    con = sqlite3.connect(DB_PATH)
    is_new = con.execute("SELECT 1 FROM users WHERE user_id=?", (user_id,)).fetchone() is None
    con.execute(
        "INSERT OR IGNORE INTO users (user_id, username, plan, req_today, req_date, referred_by) "
        "VALUES (?,?,'free',0,?,?)",
        (user_id, username, str(datetime.now().date()), referred_by)
    )
    con.execute("UPDATE users SET username=? WHERE user_id=?", (username, user_id))
    con.commit()
    con.close()
    return is_new


def set_user_field(user_id: int, field: str, value):
    if field not in ("ds_key", "ds_model"):
        raise ValueError("forbidden field")
    con = sqlite3.connect(DB_PATH)
    con.execute(f"UPDATE users SET {field}=? WHERE user_id=?", (value, user_id))
    con.commit()
    con.close()


# ─── История диалога ─────────────────────────────────────────────────────────
def add_history(user_id: int, role: str, content: str):
    content = (content or "")[:MAX_HIST_CHARS]
    con = sqlite3.connect(DB_PATH)
    con.execute(
        "INSERT INTO history (user_id, role, content, ts) VALUES (?,?,?,?)",
        (user_id, role, content, datetime.now().isoformat()),
    )
    # оставляем только последние MAX_HISTORY сообщений пользователя
    con.execute(
        "DELETE FROM history WHERE user_id=? AND id NOT IN "
        "(SELECT id FROM history WHERE user_id=? ORDER BY id DESC LIMIT ?)",
        (user_id, user_id, MAX_HISTORY),
    )
    con.commit()
    con.close()


def get_history(user_id: int) -> list:
    con = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "SELECT role, content FROM history WHERE user_id=? ORDER BY id ASC LIMIT ?",
        (user_id, MAX_HISTORY),
    ).fetchall()
    con.close()
    msgs = [{"role": r, "content": c} for r, c in rows]
    # Anthropic-формат требует, чтобы первым шёл user — срезаем ведущие assistant
    while msgs and msgs[0]["role"] != "user":
        msgs.pop(0)
    return msgs


def clear_history(user_id: int):
    con = sqlite3.connect(DB_PATH)
    con.execute("DELETE FROM history WHERE user_id=?", (user_id,))
    con.commit()
    con.close()


def is_subscribed(user) -> bool:
    if not user or not user.get("sub_until") or user.get("plan") == "free":
        return False
    return datetime.fromisoformat(user["sub_until"]) > datetime.now()


def check_and_inc_free(user_id: int) -> bool:
    con = sqlite3.connect(DB_PATH)
    row = con.execute("SELECT req_today, req_date FROM users WHERE user_id=?", (user_id,)).fetchone()
    today = str(datetime.now().date())
    if not row:
        con.close()
        return False
    req_today, req_date = row
    if req_date != today:
        req_today = 0
    if req_today >= FREE_REQUESTS_PER_DAY:
        con.close()
        return False
    con.execute("UPDATE users SET req_today=?, req_date=? WHERE user_id=?", (req_today + 1, today, user_id))
    con.commit()
    con.close()
    return True


def free_used_today(user) -> int:
    today = str(datetime.now().date())
    return user["req_today"] if user and user.get("req_date") == today else 0


def activate_plan(user_id: int, plan: str, days: int) -> str:
    con = sqlite3.connect(DB_PATH)
    row = con.execute("SELECT sub_until, plan FROM users WHERE user_id=?", (user_id,)).fetchone()
    if row and row[0] and row[1] == plan and datetime.fromisoformat(row[0]) > datetime.now():
        until = (datetime.fromisoformat(row[0]) + timedelta(days=days)).isoformat()
    else:
        until = (datetime.now() + timedelta(days=days)).isoformat()
    con.execute("UPDATE users SET sub_until=?, plan=? WHERE user_id=?", (until, plan, user_id))
    con.commit()
    con.close()
    return until


def add_referral_bonus(referrer_id: int) -> str:
    con = sqlite3.connect(DB_PATH)
    con.execute("UPDATE users SET referral_count = referral_count + 1 WHERE user_id=?", (referrer_id,))
    con.commit()
    con.close()
    user = get_user(referrer_id)
    plan = user["plan"] if user and user["plan"] != "free" else "basic"
    return activate_plan(referrer_id, plan, REFERRAL_BONUS_DAYS)


def get_stats() -> dict:
    con = sqlite3.connect(DB_PATH)
    total = con.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    now = datetime.now().isoformat()
    subs = con.execute("SELECT COUNT(*) FROM users WHERE sub_until > ? AND plan != 'free'", (now,)).fetchone()[0]
    by_plan = {}
    for plan in ["basic", "standard", "premium"]:
        by_plan[plan] = con.execute(
            "SELECT COUNT(*) FROM users WHERE plan=? AND sub_until > ?", (plan, now)
        ).fetchone()[0]
    con.close()
    return {"total": total, "subs": subs, "by_plan": by_plan}


def is_admin(user_id: int) -> bool:
    return user_id == ADMIN_ID


def get_user_plan(user) -> dict:
    if not user:
        return PLANS["free"]
    if is_subscribed(user):
        return PLANS.get(user["plan"], PLANS["free"])
    return PLANS["free"]


# ─── Форматирование текста ───────────────────────────────────────────────────
def md_to_html(text: str) -> str:
    if not text:
        return ""
    blocks = []

    def keep(s):
        blocks.append(s)
        return f"\uffff{len(blocks) - 1}\uffff"

    text = re.sub(
        r"```[ \t]*[\w+\-]*\n?(.*?)```",
        lambda m: keep("<pre><code>" + html.escape(m.group(1)) + "</code></pre>"),
        text, flags=re.DOTALL,
    )
    text = re.sub(
        r"`([^`\n]+?)`",
        lambda m: keep("<code>" + html.escape(m.group(1)) + "</code>"),
        text,
    )
    text = html.escape(text)
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s*(.+?)\s*#*$", r"<b>\1</b>", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\w)__(.+?)__(?!\w)", r"<b>\1</b>", text)
    text = re.sub(r"(?<![\*\w])\*(?=\S)(.+?)(?<=\S)\*(?![\*\w])", r"<i>\1</i>", text)
    text = re.sub(r"(?<![_\w])_(?=\S)(.+?)(?<=\S)_(?![_\w])", r"<i>\1</i>", text)
    text = re.sub(r"\[([^\]]+?)\]\((https?://[^\s)]+)\)", r'<a href="\2">\1</a>', text)
    text = re.sub(r"(?m)^\s*[-*]\s+", "• ", text)
    text = re.sub(r"\uffff(\d+)\uffff", lambda m: blocks[int(m.group(1))], text)
    return text


def split_text(text: str, limit: int = TG_LIMIT - 200) -> list:
    if len(text) <= limit:
        return [text]
    parts, cur = [], ""
    for line in text.split("\n"):
        while len(line) > limit:
            parts.append(line[:limit])
            line = line[limit:]
        if len(cur) + len(line) + 1 > limit:
            if cur:
                parts.append(cur)
            cur = line
        else:
            cur = (cur + "\n" + line) if cur else line
    if cur:
        parts.append(cur)
    return parts


def usage_bar(used: int, total: int, n: int = 5) -> str:
    filled = min(n, round((used / total) * n)) if total else 0
    return "▓" * filled + "░" * (n - filled) + f"  {used}/{total}"


# ─── Стриминг в одно/несколько сообщений ─────────────────────────────────────
class TgStreamer:
    def __init__(self, placeholder: Message):
        self.current = placeholder
        self.buffer = ""
        self.full = ""
        self.last_edit = 0.0
        self.last_render = None
        self.dirty = False
        self.last_think = 0.0
        self.think_step = 0

    @property
    def has_output(self) -> bool:
        return bool(self.full)

    async def think(self, _piece: str = None):
        """Анимация «🧠 Думаю…», пока модель размышляет и текста ещё нет."""
        if self.full:
            return
        now = time.monotonic()
        if now - self.last_think < EDIT_THROTTLE:
            return
        self.think_step = (self.think_step + 1) % 3
        dots = "." * (self.think_step + 1)
        try:
            await self.current.edit_text(f"🧠 Думаю{dots}")
        except (TelegramBadRequest, TelegramRetryAfter):
            pass
        self.last_think = now

    async def push(self, delta: str):
        self.full += delta
        if len(self.buffer) + len(delta) > SOFT_LIMIT:
            await self._flush(force=True)
            self.current = await self.current.answer("…")
            self.buffer = ""
            self.last_render = None
        self.buffer += delta
        self.dirty = True
        if time.monotonic() - self.last_edit >= EDIT_THROTTLE:
            await self._flush()

    async def _flush(self, force: bool = False):
        if not self.dirty or not self.buffer.strip():
            return
        rendered = md_to_html(self.buffer)
        if rendered == self.last_render:
            self.dirty = False
            return
        try:
            await self.current.edit_text(rendered, parse_mode="HTML")
            self.last_render = rendered
        except TelegramRetryAfter as e:
            await asyncio.sleep(e.retry_after + 0.5)
        except TelegramBadRequest:
            try:
                await self.current.edit_text(self.buffer)
                self.last_render = None
            except TelegramBadRequest:
                pass
        self.last_edit = time.monotonic()
        self.dirty = False

    async def finish(self):
        self.dirty = True
        await self._flush(force=True)

    async def fail(self, text: str):
        try:
            await self.current.edit_text(text)
        except TelegramBadRequest:
            await self.current.answer(text)


# ─── Запросы к моделям ───────────────────────────────────────────────────────
def to_gemini_contents(messages: list) -> list:
    """OpenAI-стиль (user/assistant) -> Gemini native (user/model + parts)."""
    contents = []
    for m in messages:
        role = "model" if m.get("role") == "assistant" else "user"
        contents.append({"role": role, "parts": [{"text": m.get("content", "")}]})
    return contents


async def _stream_gemini(model: str, messages: list, on_delta, api_key: str = None):
    """Родной Gemini API (работает с ключами AIza и AQ.)."""
    key = api_key or GEMINI_API_KEY
    url = f"{GEMINI_NATIVE_BASE}/models/{model}:streamGenerateContent?alt=sse"
    headers = {"x-goog-api-key": key, "Content-Type": "application/json"}
    payload = {
        "contents": to_gemini_contents(messages),
        "generationConfig": {
            "maxOutputTokens": MAX_TOKENS,
            "thinkingConfig": {"thinkingBudget": 0},  # без «размышлений» — быстрее и весь бюджет на ответ
        },
    }
    async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=20.0)) as client:
        async with client.stream("POST", url, headers=headers, json=payload) as resp:
            if resp.status_code >= 400:
                raise ApiError(resp.status_code, (await resp.aread()).decode("utf-8", "ignore"))
            async for line in resp.aiter_lines():
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if not data:
                    continue
                try:
                    obj = json.loads(data)
                    for part in obj["candidates"][0]["content"]["parts"]:
                        if part.get("text") and not part.get("thought"):
                            await on_delta(part["text"])
                except Exception:
                    continue


async def stream_completion(provider: str, model: str, messages: list, on_delta,
                            api_key: str = None) -> None:
    """Groq — OpenAI-формат; Gemini — родной API."""
    if provider == "gemini":
        await _stream_gemini(model, messages, on_delta, api_key)
        return

    url, default_key = ENDPOINTS[provider]
    key = api_key or default_key
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "max_tokens": MAX_TOKENS, "stream": True}

    async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=20.0)) as client:
        async with client.stream("POST", url, headers=headers, json=payload) as resp:
            if resp.status_code >= 400:
                body = (await resp.aread()).decode("utf-8", "ignore")
                raise ApiError(resp.status_code, body)
            async for line in resp.aiter_lines():
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if not data or data == "[DONE]":
                    continue
                try:
                    piece = json.loads(data)["choices"][0].get("delta", {}).get("content")
                    if piece:
                        await on_delta(piece)
                except Exception:
                    continue


def resolve_text_engine(user, plan: dict, admin: bool):
    """Возвращает (provider, model, keys) для текстового ответа."""
    wants_premium = admin or plan["provider"] == "gemini"
    if wants_premium:
        # игнорируем «хвосты» от старой OpenModel-версии в базе
        ds_model = user.get("ds_model") if user else None
        if not (ds_model and ds_model.startswith("gemini")):
            ds_model = None
        ds_key = user.get("ds_key") if user else None
        if not (ds_key and (ds_key.startswith("AIza") or ds_key.startswith("AQ."))):
            ds_key = None

        model = ds_model or GEMINI_PREMIUM_MODEL
        key = ds_key or GEMINI_API_KEY
        if key:
            return "gemini", model, [key]
        return "groq", GROQ_FALLBACK_MODEL, [None]   # резерв Groq, если нет ключа Gemini
    return "groq", plan["model"], [None]


async def run_text(msg: Message, messages: list, user, plan: dict, admin: bool):
    provider, model, keys = resolve_text_engine(user, plan, admin)
    placeholder = await msg.answer("✍️ Печатаю…")
    streamer = TgStreamer(placeholder)

    last_err = None
    for key in keys:
        try:
            await stream_completion(provider, model, messages, streamer.push, api_key=key)
            await streamer.finish()
            if streamer.full:
                return streamer.full
            await streamer.fail("⚠️ Модель не вернула ответ. Попробуй переформулировать.")
            return None
        except (ApiError, httpx.HTTPError) as e:
            last_err = e
            if streamer.has_output:
                await streamer.finish()
                return streamer.full
            continue

    log.error(f"AI error: {last_err}")
    if isinstance(last_err, ApiError) and last_err.status in (401, 403):
        await streamer.fail("⚠️ Ключ Gemini недействителен. Проверь его в ⚙️ Настройках "
                            "или обратись к админу.")
    elif isinstance(last_err, ApiError) and last_err.status == 429:
        await streamer.fail("⏳ Превышен бесплатный лимит Gemini (15 запросов/мин или 1500/день). "
                            "Подожди минуту и попробуй снова.")
    else:
        await streamer.fail("⚠️ Не получилось получить ответ. Попробуй ещё раз чуть позже.")
    return None


# ─── Vision (фото) — через Groq ──────────────────────────────────────────────
async def run_vision(msg: Message, caption: str, image_b64: str):
    messages = [{
        "role": "user",
        "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}},
            {"type": "text", "text": caption or "Опиши, что на изображении"},
        ],
    }]
    placeholder = await msg.answer("🖼 Смотрю на фото…")
    streamer = TgStreamer(placeholder)
    try:
        await stream_completion("groq", GROQ_VISION_MODEL, messages, streamer.push)
        await streamer.finish()
    except (ApiError, httpx.HTTPError) as e:
        log.error(f"Vision error: {e}")
        await streamer.fail("⚠️ Не удалось разобрать изображение. Попробуй другое фото.")


# ─── Извлечение текста из файла ──────────────────────────────────────────────
async def extract_text_from_file(msg: Message, bot: Bot) -> str:
    doc = msg.document
    file = await bot.get_file(doc.file_id)
    buf = io.BytesIO()
    await bot.download_file(file.file_path, buf)
    content = buf.getvalue()
    fname = (doc.file_name or "").lower()

    if doc.mime_type == "text/plain" or fname.endswith(".txt"):
        try:
            return content.decode("utf-8")[:8000]
        except UnicodeDecodeError:
            return content.decode("latin-1", "ignore")[:8000]

    if doc.mime_type == "application/pdf" or fname.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = [(p.extract_text() or "") for p in reader.pages[:20]]
            result = "\n".join(pages).strip()
            if len(result) > 30:
                return result[:8000]
        except Exception as e:
            log.error(f"PDF error: {e}")
        return "⚠️ Не удалось извлечь текст из PDF (возможно, это скан без текстового слоя)."

    DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if doc.mime_type == DOCX_MIME or fname.endswith(".docx"):
        try:
            from docx import Document
            d = Document(io.BytesIO(content))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            result = "\n".join(parts).strip()
            if len(result) > 10:
                return result[:8000]
        except Exception as e:
            log.error(f"DOCX error: {e}")
        return "⚠️ Не удалось прочитать DOCX-файл."

    if fname.endswith(".doc"):
        return "⚠️ Старый формат .doc не поддерживается. Пересохрани файл как .docx (Файл → Сохранить как)."

    return f"⚠️ Тип файла «{doc.mime_type}» не поддерживается. Поддерживаются: TXT, PDF, DOCX."


# ─── Распознавание голоса (Groq Whisper) ─────────────────────────────────────
async def transcribe_voice(content: bytes, filename: str = "voice.ogg") -> str:
    url = "https://api.groq.com/openai/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files = {"file": (filename, content, "audio/ogg")}
    data = {"model": GROQ_WHISPER_MODEL, "response_format": "json"}
    async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=20.0)) as client:
        resp = await client.post(url, headers=headers, files=files, data=data)
        if resp.status_code >= 400:
            raise ApiError(resp.status_code, resp.text)
        return (resp.json().get("text") or "").strip()


# ─── Клавиатуры ──────────────────────────────────────────────────────────────
def kb_main(user_id: int) -> InlineKeyboardMarkup:
    buttons = []
    if is_admin(user_id):
        buttons.append([InlineKeyboardButton(text="👑 Админ-панель", callback_data="admin_panel")])
    buttons.extend([
        [InlineKeyboardButton(text="💎 Тарифы и подписка", callback_data="plans")],
        [InlineKeyboardButton(text="⚙️ Настройки", callback_data="settings")],
        [InlineKeyboardButton(text="🧹 Очистить диалог", callback_data="reset_dialog")],
        [InlineKeyboardButton(text="👥 Реферальная программа", callback_data="referral")],
        [InlineKeyboardButton(text="📊 Мой статус", callback_data="status")],
    ])
    return InlineKeyboardMarkup(inline_keyboard=buttons)


def kb_plans() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⚡ Базовый — 150 ⭐/мес", callback_data="buy_basic")],
        [InlineKeyboardButton(text="🔥 Стандарт — 250 ⭐/мес", callback_data="buy_standard")],
        [InlineKeyboardButton(text="👑 Премиум — 450 ⭐/мес", callback_data="buy_premium")],
        [InlineKeyboardButton(text="◀️ Назад", callback_data="back_main")],
    ])


def kb_back() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="◀️ Назад", callback_data="back_main")]
    ])


def kb_admin() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Статистика", callback_data="admin_stats")],
        [InlineKeyboardButton(text="◀️ Назад", callback_data="back_main")],
    ])


def kb_settings(user, admin: bool) -> InlineKeyboardMarkup:
    rows = []
    can_premium = admin or (is_subscribed(user) and user.get("plan") == "premium")
    if can_premium:
        cur_model = (user.get("ds_model") or GEMINI_PREMIUM_MODEL) if user else GEMINI_PREMIUM_MODEL
        flash_mark = "✅" if cur_model == GEMINI_PREMIUM_MODEL else "▫️"
        lite_mark = "✅" if cur_model == GEMINI_FAST_MODEL else "▫️"
        rows.append([
            InlineKeyboardButton(text=f"{flash_mark} Flash (умнее)", callback_data="model_fast"),
            InlineKeyboardButton(text=f"{lite_mark} Flash-Lite (быстрее)", callback_data="model_pro"),
        ])
        if user and user.get("ds_key"):
            rows.append([InlineKeyboardButton(text="🔑 Заменить мой ключ", callback_data="set_ds_key")])
            rows.append([InlineKeyboardButton(text="🗑 Удалить мой ключ", callback_data="del_ds_key")])
        else:
            rows.append([InlineKeyboardButton(text="🔑 Вставить свой Gemini-ключ", callback_data="set_ds_key")])
    rows.append([InlineKeyboardButton(text="◀️ Назад", callback_data="back_main")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


# ─── FSM ─────────────────────────────────────────────────────────────────────
class SettingsSG(StatesGroup):
    waiting_key = State()


def welcome_text(name: str) -> str:
    return (
        f"👋 Привет, <b>{html.escape(name)}</b>!\n\n"
        f"Я — AI-ассистент с выбором модели 🤖\n"
        f"━━━━━━━━━━━━━━━\n"
        f"🆓 <b>Бесплатно</b> — 3 запроса/день\n"
        f"⚡ <b>Базовый</b> — безлимит + файлы\n"
        f"🔥 <b>Стандарт</b> — + анализ фото\n"
        f"👑 <b>Премиум</b> — Gemini 2.5 Flash, всё включено\n"
        f"━━━━━━━━━━━━━━━\n"
        f"Просто напиши мне что-нибудь — отвечу 👇"
    )


dp = Dispatcher(storage=MemoryStorage())


@dp.message(CommandStart())
async def cmd_start(msg: Message, state: FSMContext):
    await state.clear()
    user_id = msg.from_user.id
    username = msg.from_user.username or ""
    referred_by = None
    args = msg.text.split()
    if len(args) > 1 and args[1].startswith("ref_"):
        try:
            ref_id = int(args[1].split("_")[1])
            if ref_id != user_id:
                referred_by = ref_id
        except (ValueError, IndexError):
            pass

    is_new = upsert_user(user_id, username, referred_by)

    if is_new and referred_by and get_user(referred_by):
        until = add_referral_bonus(referred_by)
        until_str = datetime.fromisoformat(until).strftime("%d.%m.%Y")
        try:
            await msg.bot.send_message(
                referred_by,
                f"🎉 По твоей ссылке пришёл новый пользователь!\n"
                f"✅ Начислено <b>+{REFERRAL_BONUS_DAYS} дней</b> подписки (до {until_str})",
                parse_mode="HTML",
            )
        except Exception:
            pass

    await msg.answer(welcome_text(msg.from_user.first_name), parse_mode="HTML",
                     reply_markup=kb_main(user_id))


@dp.callback_query(F.data == "plans")
async def cb_plans(cb: CallbackQuery):
    await cb.message.edit_text(
        "💎 <b>Тарифы</b>\n"
        "━━━━━━━━━━━━━━━\n"
        "⚡ <b>Базовый</b> — 150 ⭐/мес\n  Llama 3.3 70B · безлимит · файлы\n\n"
        "🔥 <b>Стандарт</b> — 250 ⭐/мес\n  + анализ фото\n\n"
        "👑 <b>Премиум</b> — 450 ⭐/мес\n  Gemini 2.5 Flash · топ-качество\n"
        "  можно подключить свой Gemini-ключ\n"
        "━━━━━━━━━━━━━━━\n"
        "Оплата — звёздами Telegram ⭐",
        parse_mode="HTML", reply_markup=kb_plans(),
    )
    await cb.answer()


@dp.callback_query(F.data.startswith("buy_"))
async def cb_buy(cb: CallbackQuery, bot: Bot):
    plan_key = cb.data.replace("buy_", "")
    plan = PLANS.get(plan_key)
    if not plan or plan_key == "free":
        await cb.answer("❌ Тариф не найден", show_alert=True)
        return
    await bot.send_invoice(
        chat_id=cb.from_user.id,
        title=f"{plan['name']} — AI подписка",
        description=plan["desc"],
        payload=f"plan_{plan_key}",
        currency="XTR",
        prices=[LabeledPrice(label=plan["name"], amount=plan["price_stars"])],
    )
    await cb.answer()


@dp.pre_checkout_query()
async def pre_checkout(query: PreCheckoutQuery, bot: Bot):
    await bot.answer_pre_checkout_query(query.id, ok=True)


@dp.message(F.successful_payment)
async def payment_success(msg: Message):
    payload = msg.successful_payment.invoice_payload
    plan_key = payload.replace("plan_", "")
    plan = PLANS.get(plan_key, PLANS["basic"])
    until = activate_plan(msg.from_user.id, plan_key, SUBSCRIPTION_DAYS)
    until_str = datetime.fromisoformat(until).strftime("%d.%m.%Y")
    await msg.answer(
        f"🎉 <b>Тариф {plan['name']} активирован!</b>\n"
        f"Действует до {until_str}\n\nМодель: <b>{plan['model']}</b>\n{plan['desc']}",
        parse_mode="HTML", reply_markup=kb_main(msg.from_user.id),
    )


@dp.callback_query(F.data == "status")
async def cb_status(cb: CallbackQuery):
    user_id = cb.from_user.id
    user = get_user(user_id)
    if not user:
        upsert_user(user_id, cb.from_user.username or "")
        user = get_user(user_id)

    plan = get_user_plan(user)
    if is_admin(user_id):
        status = "👑 <b>Администратор</b>\n♾️ Безлимит · Gemini 2.5 Flash"
    elif is_subscribed(user):
        until = datetime.fromisoformat(user["sub_until"]).strftime("%d.%m.%Y")
        status = f"{plan['name']}\nДо: <b>{until}</b>\nМодель: <code>{plan['model']}</code>"
    else:
        used = free_used_today(user)
        status = (f"🆓 <b>Бесплатный план</b>\n"
                  f"Лимит на сегодня:\n<code>{usage_bar(used, FREE_REQUESTS_PER_DAY)}</code>")

    await cb.message.edit_text(
        f"📊 <b>Твой статус</b>\n━━━━━━━━━━━━━━━\n{status}\n━━━━━━━━━━━━━━━\n"
        f"👥 Приглашено друзей: <b>{user['referral_count']}</b>",
        parse_mode="HTML", reply_markup=kb_main(user_id),
    )
    await cb.answer()


@dp.callback_query(F.data == "settings")
async def cb_settings(cb: CallbackQuery):
    user_id = cb.from_user.id
    user = get_user(user_id)
    if not user:
        upsert_user(user_id, cb.from_user.username or "")
        user = get_user(user_id)
    admin = is_admin(user_id)
    can_premium = admin or (is_subscribed(user) and user.get("plan") == "premium")

    if not can_premium:
        await cb.message.edit_text(
            "⚙️ <b>Настройки</b>\n━━━━━━━━━━━━━━━\n"
            "Выбор модели и подключение своего Gemini-ключа доступны "
            "на тарифе 👑 <b>Премиум</b>.",
            parse_mode="HTML", reply_markup=kb_back(),
        )
        await cb.answer()
        return

    cur_model = (user.get("ds_model") or GEMINI_PREMIUM_MODEL)
    key_state = "свой ключ подключён 🔑" if user.get("ds_key") else "используются ключи бота"
    await cb.message.edit_text(
        "⚙️ <b>Настройки (Премиум)</b>\n━━━━━━━━━━━━━━━\n"
        f"Текущая модель: <code>{cur_model}</code>\n"
        f"Ключ: {key_state}\n━━━━━━━━━━━━━━━\n"
        "• <b>Flash</b> — Gemini 2.5 Flash, умнее\n"
        "• <b>Flash-Lite</b> — быстрее и легче по лимитам\n\n"
        "Можешь вставить свой Gemini-ключ (AIza…) — тогда запросы пойдут через него.",
        parse_mode="HTML", reply_markup=kb_settings(user, admin),
    )
    await cb.answer()


@dp.callback_query(F.data == "model_fast")
async def cb_model_fast(cb: CallbackQuery):
    set_user_field(cb.from_user.id, "ds_model", GEMINI_PREMIUM_MODEL)
    await cb.answer("Модель: Flash (умнее) ✅")
    await cb_settings(cb)


@dp.callback_query(F.data == "model_pro")
async def cb_model_pro(cb: CallbackQuery):
    set_user_field(cb.from_user.id, "ds_model", GEMINI_FAST_MODEL)
    await cb.answer("Модель: Flash-Lite (быстрее) ✅")
    await cb_settings(cb)


@dp.callback_query(F.data == "del_ds_key")
async def cb_del_key(cb: CallbackQuery):
    set_user_field(cb.from_user.id, "ds_key", None)
    await cb.answer("Ключ удалён, вернулись на ключи бота ✅")
    await cb_settings(cb)


@dp.callback_query(F.data == "set_ds_key")
async def cb_set_key(cb: CallbackQuery, state: FSMContext):
    user = get_user(cb.from_user.id)
    admin = is_admin(cb.from_user.id)
    can_premium = admin or (is_subscribed(user) and user.get("plan") == "premium")
    if not can_premium:
        await cb.answer("Доступно на Премиуме", show_alert=True)
        return
    await state.set_state(SettingsSG.waiting_key)
    await cb.message.answer(
        "🔑 Пришли свой ключ Gemini одним сообщением (начинается с <code>AIza</code> или <code>AQ.</code>).\n"
        "Получить бесплатно (без карты) на aistudio.google.com.\n\nЧтобы отменить — отправь /cancel",
        parse_mode="HTML",
    )
    await cb.answer()


@dp.message(Command("cancel"))
async def cmd_cancel(msg: Message, state: FSMContext):
    if await state.get_state():
        await state.clear()
        await msg.answer("Отменено.", reply_markup=kb_main(msg.from_user.id))


@dp.message(Command("reset"))
async def cmd_reset(msg: Message):
    clear_history(msg.from_user.id)
    await msg.answer("🧹 История диалога очищена. Начинаем с чистого листа.")


@dp.callback_query(F.data == "reset_dialog")
async def cb_reset(cb: CallbackQuery):
    clear_history(cb.from_user.id)
    await cb.answer("🧹 Диалог очищен", show_alert=True)


@dp.message(StateFilter(SettingsSG.waiting_key), F.text)
async def receive_ds_key(msg: Message, state: FSMContext):
    key = msg.text.strip()
    await state.clear()
    try:
        await msg.delete()
    except TelegramBadRequest:
        pass
    if not ((key.startswith("AIza") or key.startswith("AQ.")) and len(key) >= 20):
        await msg.answer("❌ Это не похоже на ключ Gemini (должен начинаться с AIza или AQ.). "
                         "Попробуй ещё раз через ⚙️ Настройки.")
        return
    set_user_field(msg.from_user.id, "ds_key", key)
    await msg.answer("✅ Ключ сохранён. Теперь Премиум-запросы идут через твой ключ.\n"
                     "(Сообщение с ключом я удалил из чата.)",
                     reply_markup=kb_main(msg.from_user.id))


@dp.callback_query(F.data == "referral")
async def cb_referral(cb: CallbackQuery):
    user = get_user(cb.from_user.id) or {"referral_count": 0}
    ref_link = f"https://t.me/{BOT_USERNAME}?start=ref_{cb.from_user.id}"
    await cb.message.edit_text(
        f"👥 <b>Реферальная программа</b>\n━━━━━━━━━━━━━━━\n"
        f"За каждого друга: <b>+{REFERRAL_BONUS_DAYS} дней</b> подписки\n"
        f"Приглашено: <b>{user['referral_count']} чел.</b>\n━━━━━━━━━━━━━━━\n"
        f"Твоя ссылка:\n<code>{ref_link}</code>",
        parse_mode="HTML", reply_markup=kb_back(),
    )
    await cb.answer()


@dp.callback_query(F.data == "back_main")
async def cb_back(cb: CallbackQuery, state: FSMContext):
    await state.clear()
    await cb.message.edit_text("🏠 Главное меню", reply_markup=kb_main(cb.from_user.id))
    await cb.answer()


@dp.callback_query(F.data == "admin_panel")
async def cb_admin_panel(cb: CallbackQuery):
    if not is_admin(cb.from_user.id):
        await cb.answer("⛔ Нет доступа", show_alert=True)
        return
    stats = get_stats()
    keys_note = "задан ✅" if GEMINI_API_KEY else "не задан ⚠️"
    await cb.message.edit_text(
        f"👑 <b>Админ-панель</b>\n━━━━━━━━━━━━━━━\n"
        f"👥 Всего пользователей: <b>{stats['total']}</b>\n"
        f"💎 Активных подписок: <b>{stats['subs']}</b>\n\n"
        f"⚡ Базовых: <b>{stats['by_plan']['basic']}</b>\n"
        f"🔥 Стандарт: <b>{stats['by_plan']['standard']}</b>\n"
        f"👑 Премиум: <b>{stats['by_plan']['premium']}</b>\n\n"
        f"🔑 Ключ Gemini: <b>{keys_note}</b>\n━━━━━━━━━━━━━━━\n"
        f"Выдать подписку:\n<code>/give USER_ID план дней</code>\n"
        f"Планы: basic, standard, premium\n"
        f"Пример: <code>/give 123456789 premium 30</code>",
        parse_mode="HTML", reply_markup=kb_admin(),
    )
    await cb.answer()


@dp.callback_query(F.data == "admin_stats")
async def cb_admin_stats(cb: CallbackQuery):
    if not is_admin(cb.from_user.id):
        await cb.answer("⛔ Нет доступа", show_alert=True)
        return
    stats = get_stats()
    await cb.answer(f"👥 Пользователей: {stats['total']}\n💎 Подписок: {stats['subs']}", show_alert=True)


@dp.message(Command("give"))
async def cmd_give(msg: Message):
    if not is_admin(msg.from_user.id):
        return
    args = msg.text.split()
    if len(args) != 4:
        await msg.answer("Использование: /give [user_id] [план] [дней]\nПример: /give 123456789 premium 30")
        return
    try:
        target_id, plan_key, days = int(args[1]), args[2], int(args[3])
    except ValueError:
        await msg.answer("❌ Неверный формат.")
        return
    if plan_key not in PLANS or plan_key == "free":
        await msg.answer("❌ Неверный план. Доступны: basic, standard, premium")
        return
    if not get_user(target_id):
        await msg.answer(f"❌ Пользователь {target_id} не найден.")
        return
    until = activate_plan(target_id, plan_key, days)
    until_str = datetime.fromisoformat(until).strftime("%d.%m.%Y")
    plan = PLANS[plan_key]
    await msg.answer(f"✅ Пользователю {target_id} выдан тариф {plan['name']} до {until_str}")
    try:
        await msg.bot.send_message(target_id, f"🎁 Тебе выдан тариф {plan['name']} до {until_str}!")
    except Exception:
        pass


@dp.message(F.photo)
async def handle_photo(msg: Message, bot: Bot):
    user_id = msg.from_user.id
    upsert_user(user_id, msg.from_user.username or "")
    user = get_user(user_id)
    admin = is_admin(user_id)
    plan = PLANS["premium"] if admin else get_user_plan(user)

    if not admin and not is_subscribed(user):
        await msg.answer("📸 Анализ фото доступен на платных тарифах.\n\nОформи подписку 👇",
                         reply_markup=kb_plans())
        return
    if not admin and not plan["vision"]:
        await msg.answer(f"📸 Фото недоступно на тарифе {plan['name']}.\nНужен 🔥 Стандарт или 👑 Премиум.",
                         reply_markup=kb_plans())
        return

    await bot.send_chat_action(msg.chat.id, "typing")
    photo = msg.photo[-1]
    file = await bot.get_file(photo.file_id)
    buf = io.BytesIO()
    await bot.download_file(file.file_path, buf)
    image_b64 = base64.b64encode(buf.getvalue()).decode()
    await run_vision(msg, msg.caption or "Опиши, что на изображении", image_b64)


@dp.message(F.document)
async def handle_document(msg: Message, bot: Bot):
    user_id = msg.from_user.id
    upsert_user(user_id, msg.from_user.username or "")
    user = get_user(user_id)
    admin = is_admin(user_id)
    plan = PLANS["premium"] if admin else get_user_plan(user)

    if not admin and not is_subscribed(user):
        await msg.answer("📄 Работа с файлами доступна на платных тарифах.\n\nОформи подписку 👇",
                         reply_markup=kb_plans())
        return
    if not admin and not plan["files"]:
        await msg.answer(f"📄 Файлы недоступны на тарифе {plan['name']}.", reply_markup=kb_plans())
        return

    await bot.send_chat_action(msg.chat.id, "typing")
    text = await extract_text_from_file(msg, bot)
    if text.startswith("⚠️"):
        await msg.answer(text)
        return

    caption = msg.caption or "Проанализируй этот текст и кратко изложи суть."
    messages = [{"role": "user", "content": f"{caption}\n\n---\n{text}"}]
    await run_text(msg, messages, user, plan, admin)


@dp.message(F.voice | F.audio)
async def handle_voice(msg: Message, bot: Bot):
    user_id = msg.from_user.id
    upsert_user(user_id, msg.from_user.username or "")
    user = get_user(user_id)
    admin = is_admin(user_id)

    if not admin:
        plan = get_user_plan(user)
        if not is_subscribed(user) and not check_and_inc_free(user_id):
            await msg.answer(
                f"⛔ Бесплатный лимит исчерпан ({FREE_REQUESTS_PER_DAY} запроса/день).\n\n"
                f"Выбери тариф для продолжения 👇",
                reply_markup=kb_plans(),
            )
            return
    else:
        plan = PLANS["premium"]

    media = msg.voice or msg.audio
    status = await msg.answer("🎤 Распознаю голос…")
    try:
        file = await bot.get_file(media.file_id)
        buf = io.BytesIO()
        await bot.download_file(file.file_path, buf)
        text = await transcribe_voice(buf.getvalue(),
                                      getattr(media, "file_name", None) or "voice.ogg")
    except (ApiError, httpx.HTTPError) as e:
        log.error(f"STT error: {e}")
        await status.edit_text("⚠️ Не удалось распознать голос. Попробуй ещё раз.")
        return

    if not text:
        await status.edit_text("🤷 Не расслышал. Запиши ещё раз почётче.")
        return

    await status.edit_text(f"🎤 <i>Распознал:</i> {html.escape(text)}", parse_mode="HTML")
    history = get_history(user_id)
    messages = history + [{"role": "user", "content": text}]
    reply = await run_text(msg, messages, user, plan, admin)
    if reply:
        add_history(user_id, "user", text)
        add_history(user_id, "assistant", reply)


@dp.message(StateFilter(None), F.text & ~F.text.startswith("/"))
async def handle_message(msg: Message):
    user_id = msg.from_user.id
    upsert_user(user_id, msg.from_user.username or "")
    user = get_user(user_id)
    admin = is_admin(user_id)

    if not admin:
        plan = get_user_plan(user)
        if not is_subscribed(user):
            if not check_and_inc_free(user_id):
                await msg.answer(
                    f"⛔ Бесплатный лимит исчерпан ({FREE_REQUESTS_PER_DAY} запроса/день).\n\n"
                    f"Выбери тариф для продолжения 👇",
                    reply_markup=kb_plans(),
                )
                return
    else:
        plan = PLANS["premium"]

    await msg.bot.send_chat_action(msg.chat.id, "typing")
    history = get_history(user_id)
    messages = history + [{"role": "user", "content": msg.text}]
    reply = await run_text(msg, messages, user, plan, admin)
    if reply:
        add_history(user_id, "user", msg.text)
        add_history(user_id, "assistant", reply)


async def main():
    init_db()
    if not GEMINI_API_KEY:
        log.warning("GEMINI_API_KEY не задан — Премиум будет отвечать на резервной модели Groq.")
    bot = Bot(token=BOT_TOKEN)
    log.info("Бот запущен ✅")
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
